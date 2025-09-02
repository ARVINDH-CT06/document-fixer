from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL

def format_document(doc):
    changes_report = []

    # 0. Page setup: margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    changes_report.append("Set page margins to 1 inch on all sides")

    # 1. Format headers & footers
    header_count, footer_count = 0, 0
    for section in doc.sections:
        header = section.header
        footer = section.footer
        for para in header.paragraphs:
            for run in para.runs:
                run.font.name = "Times New Roman"
                run.font.size = Pt(12)
            header_count += 1
        for para in footer.paragraphs:
            for run in para.runs:
                run.font.name = "Times New Roman"
                run.font.size = Pt(12)
            footer_count += 1
    changes_report.append(f"Formatted {header_count} header and {footer_count} footer paragraphs")

    # 2. Format paragraphs + bullets
    para_count, bullet_count = 0, 0
    for para in doc.paragraphs:
        if para.text.strip():
            # Apply default paragraph formatting
            for run in para.runs:
                run.font.name = "Times New Roman"
                run.font.size = Pt(12)
            para.paragraph_format.line_spacing = 1.5
            para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            para_count += 1

            # Detect bullets or numbering
            if para.style.name in ["List Bullet", "List Number"]:
                for run in para.runs:
                    run.font.name = "Times New Roman"
                    run.font.size = Pt(12)
                bullet_count += 1
    changes_report.append(f"Formatted {para_count} paragraphs (Times New Roman, 12pt, justified)")
    changes_report.append(f"Fixed {bullet_count} bullet/numbered list items")

    # 3. Format headings
    heading_count = 0
    for para in doc.paragraphs:
        text = para.text.strip()
        if text.isupper() and len(text.split()) <= 6:
            for run in para.runs:
                run.font.size = Pt(16)
                run.font.bold = True
                run.font.color.rgb = RGBColor(0, 0, 255)
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            heading_count += 1
        elif text.istitle():
            for run in para.runs:
                run.font.size = Pt(14)
                run.font.bold = True
                run.font.color.rgb = RGBColor(0, 0, 255)
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            heading_count += 1
    changes_report.append(f"Formatted {heading_count} headings (center aligned, bold, blue)")

    # 4. Format tables
    table_count = 0
    for table in doc.tables:
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        for row_idx, row in enumerate(table.rows):
            for cell in row.cells:
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                for para in cell.paragraphs:
                    if row_idx == 0:  # header row
                        for run in para.runs:
                            run.font.bold = True
                            run.font.name = "Times New Roman"
                            run.font.size = Pt(12)
                        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    else:
                        if para.text.strip().isdigit():
                            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        elif len(para.text) > 30:
                            para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                        else:
                            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        table_count += 1
    changes_report.append(f"Formatted {table_count} tables (center aligned, bold headers)")

    return doc, changes_report
